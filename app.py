import streamlit as st
import ollama
import os
from datetime import datetime
import psutil
import PyPDF2
from docx import Document

# ========================================
# CONFIGURATION
# ========================================
MAX_MESSAGES = 20
PRESERVE_SYSTEM = True
TEMP_DIR = "temp_uploads"

# ========================================
# UTILITY FUNCTIONS
# ========================================

def prune_conversation(messages, max_messages=MAX_MESSAGES):
    """Keep only recent messages to prevent token overflow"""
    if len(messages) <= max_messages:
        return messages
    
    if PRESERVE_SYSTEM and messages[0]["role"] == "system":
        return [messages[0]] + messages[-(max_messages-1):]
    
    return messages[-max_messages:]

# ========================================
# TOOL FUNCTIONS
# ========================================

def get_system_stats():
    """Returns real-time CPU and RAM usage percentage."""
    cpu = psutil.cpu_percent(interval=1)
    ram = psutil.virtual_memory().percent
    return f"CPU: {cpu}% | RAM: {ram}%"

def list_files(directory_path="."):
    """Liste les fichiers d'un dossier."""
    try:
        if directory_path in ["", ".", "./"]: 
            directory_path = os.getcwd()
        files = os.listdir(directory_path)
        return f"Fichiers dans {directory_path}: " + ", ".join(files)
    except Exception as e:
        return f"Erreur : {str(e)}"

def summarize_directory(directory_path: str = "."):
    """Lit et résume les fichiers d'un dossier (txt, py, pdf, docx)."""
    if directory_path in ["", ".", "./"]:
        directory_path = os.getcwd()

    if not os.path.exists(directory_path):
        return f"Erreur : Le chemin '{directory_path}' n'existe pas."

    combined_text = ""
    files_processed = []

    try:
        for filename in os.listdir(directory_path):
            file_path = os.path.join(directory_path, filename)

            if filename.endswith(('.txt', '.md', '.py')):
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        combined_text += f"\n[Fichier: {filename}]\n{f.read()[:1500]}\n"
                        files_processed.append(filename)
                except: 
                    continue

            elif filename.endswith('.pdf'):
                try:
                    with open(file_path, 'rb') as f:
                        reader = PyPDF2.PdfReader(f)
                        pdf_text = "".join([page.extract_text() or "" for page in reader.pages[:3]])
                        combined_text += f"\n[Fichier PDF: {filename}]\n{pdf_text[:1500]}\n"
                        files_processed.append(filename)
                except: 
                    continue

            elif filename.endswith('.docx'):
                try:
                    doc = Document(file_path)
                    full_text = [para.text for para in doc.paragraphs]
                    docx_content = "\n".join(full_text)[:1500]
                    combined_text += f"\n[Fichier Word: {filename}]\n{docx_content}\n"
                    files_processed.append(filename)
                except: 
                    continue

    except Exception as e:
        return f"Erreur d'accès : {str(e)}"

    if not combined_text:
        return "Aucun document (txt, py, pdf, docx) trouvé."

    prompt = f"Fais un résumé structuré des documents suivants :\n\n{combined_text}"

    try:
        response = ollama.chat(model="llama3.2", messages=[{"role": "user", "content": prompt}])
        return f"Analyse de : {', '.join(files_processed)}\n\n{response['message']['content']}"
    except Exception as e:
        return f"Erreur Ollama : {str(e)}"

def read_specific_file(file_path: str):
    """Lit le contenu d'un fichier spécifique."""
    if not os.path.exists(file_path):
        return f"Erreur : Le fichier '{file_path}' est introuvable."

    try:
        if file_path.endswith(('.txt', '.md', '.py')):
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()[:3000]

        elif file_path.endswith('.pdf'):
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                return "".join([p.extract_text() for p in reader.pages[:5]])[:3000]

        elif file_path.endswith('.docx'):
            doc = Document(file_path)
            return "\n".join([p.text for p in doc.paragraphs])[:3000]

        return "Format de fichier non supporté."
    except Exception as e:
        return f"Erreur lors de la lecture : {str(e)}"

def triage_emails(max_results=5):
    """Trie et affiche les emails importants."""
    return (
        "Liste des emails triés :\n"
        "[Urgent] Facture EDF - À payer avant le 10 février.\n"
        "[Info] Réunion ENSIAS - Présentation IoT lundi prochain."
    )

def manage_tasks(action: str, task_text: str = ""):
    """Actions: 'add' pour sauvegarder, 'list' pour afficher les rappels."""
    filename = "reminders.txt"
    if action == "add":
        with open(filename, "a") as f:
            f.write(f"[{datetime.now().strftime('%Y-%m-%d %H:%M')}] {task_text}\n")
        return f"Rappel sauvegardé : {task_text}"
    elif action == "list":
        if not os.path.exists(filename): 
            return "Aucun rappel trouvé."
        with open(filename, "r") as f:
            return f.read()

# ========================================
# STREAMLIT APP
# ========================================

st.set_page_config(page_title="AI Agent OS Supervisor", page_icon="🤖", layout="wide")

st.title("🤖 Agent TalentAI-Bot")
st.markdown("Interface de supervision et d'analyse de documents (PDF, Word, Texte)")

# ========================================
# SIDEBAR
# ========================================

with st.sidebar:
    # PR #3: File Upload
    st.header("📁 Upload de Documents")
    
    uploaded_files = st.file_uploader(
        "Déposez vos fichiers ici",
        accept_multiple_files=True,
        type=['pdf', 'txt', 'docx', 'doc', 'md', 'py'],
        help="Les fichiers seront analysables immédiatement"
    )
    
    if uploaded_files:
        os.makedirs(TEMP_DIR, exist_ok=True)
        
        saved_files = []
        for uploaded_file in uploaded_files:
            file_path = os.path.join(TEMP_DIR, uploaded_file.name)
            with open(file_path, "wb") as f:
                f.write(uploaded_file.getbuffer())
            saved_files.append(uploaded_file.name)
        
        st.success(f"✅ {len(saved_files)} fichier(s) uploadé(s)")
        with st.expander("Fichiers disponibles"):
            for fname in saved_files:
                st.text(f"📄 {fname}")
    
    st.divider()
    
    # System Stats
    st.header("Système")
    if st.button("Actualiser les stats"):
        stats = get_system_stats()
        st.write(stats)
    
    # PR #2: Conversation Management
    st.divider()
    col1, col2 = st.columns(2)
    with col1:
        msg_count = len([m for m in st.session_state.get('messages', []) if m.get("role") != "system"])
        st.metric("Messages", msg_count)
    with col2:
        if st.button("🗑️ Effacer", help="Réinitialiser la conversation"):
            st.session_state.messages = [st.session_state.messages[0]]
            st.rerun()
    
    st.divider()
    st.info("Agents : Supervisor, Summarizer, Email Triage, Task Manager")

# ========================================
# SESSION STATE INITIALIZATION
# ========================================

if "messages" not in st.session_state:
    st.session_state.messages = [
        {
            "role": "system", 
            "content": (
                "Tu es un assistant IA polyvalent nommé TalentAI-Bot. "
                "Tu peux fonctionner comme un Chatbot ou comme un Agent. "
                f"Les fichiers uploadés se trouvent dans '{TEMP_DIR}/'. "
                "Tu as accès à 'manage_tasks' pour gérer les rappels. "
                "Si l'utilisateur mentionne 'mails', 'emails' ou 'courriels', "
                "tu dois EXCLUSIVEMENT utiliser l'outil 'triage_emails'. "
                "N'utilise jamais 'list_files' ou 'summarize_directory' pour les emails."
            )
        }
    ]

# ========================================
# CHAT DISPLAY
# ========================================

for message in st.session_state.messages:
    if message["role"] != "system":
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

# ========================================
# CHAT INPUT & PROCESSING
# ========================================

if prompt := st.chat_input("Comment puis-je vous aider ?"):
    st.session_state.messages.append({"role": "user", "content": prompt})
    with st.chat_message("user"):
        st.markdown(prompt)

    with st.chat_message("assistant"):
        available_tools = [
            get_system_stats, 
            list_files, 
            summarize_directory, 
            read_specific_file, 
            triage_emails,
            manage_tasks
        ]

        # PR #2: Use pruned messages
        response = ollama.chat(
            model="llama3.2", 
            messages=prune_conversation(st.session_state.messages), 
            tools=available_tools
        )

        if response.get('message', {}).get('tool_calls'):
            st.session_state.messages.append(response['message'])

            # PR #1: Real-time Tool Execution Feedback
            for tool in response['message']['tool_calls']:
                name = tool['function']['name']
                args = tool['function']['arguments']
                
                with st.status(f"🔧 Exécution: {name}...", expanded=True) as status:
                    st.write(f"**Arguments:** `{args}`")
                    
                    if name == "get_system_stats": 
                        result = get_system_stats()
                    elif name == "list_files": 
                        result = list_files(**args)
                    elif name == "summarize_directory": 
                        result = summarize_directory(**args)
                    elif name == "read_specific_file": 
                        result = read_specific_file(**args)
                    elif name == "triage_emails": 
                        result = triage_emails(**args)
                    elif name == "manage_tasks": 
                        result = manage_tasks(**args)
                    else: 
                        result = f"Outil '{name}' inconnu."
                    
                    st.write("**Résultat:**")
                    preview = str(result)[:500] + ("..." if len(str(result)) > 500 else "")
                    st.code(preview, language="text")
                    status.update(label=f"✅ {name} terminé", state="complete")
                
                st.session_state.messages.append({'role': 'tool', 'content': result})

            final_res = ollama.chat(model="llama3.2", messages=prune_conversation(st.session_state.messages))
            full_response = final_res['message']['content']
        else:
            full_response = response['message']['content']

        st.markdown(full_response)
        st.session_state.messages.append({"role": "assistant", "content": full_response})import streamlit as st
import ollama
import os
from datetime import datetime
import psutil
import PyPDF2
from docx import Document

# ========================================
# CONFIGURATION
# ========================================
MAX_MESSAGES = 20
PRESERVE_SYSTEM = True
TEMP_DIR = "temp_uploads"

# ========================================
# UTILITY FUNCTIONS
# ========================================

def prune_conversation(messages, max_messages=MAX_MESSAGES):
    """Keep only recent messages to prevent token overflow"""
    if len(messages) <= max_messages:
        return messages
    
    if PRESERVE_SYSTEM and messages[0]["role"] == "system":
        return [messages[0]] + messages[-(max_messages-1):]
    
    return messages[-max_messages:]

# ========================================
# TOOL FUNCTIONS
# ========================================

def get_system_stats():
    """Returns real-time CPU and RAM usage percentage."""
    cpu = psutil.cpu_percent(interval=1)
    ram = psutil.virtual_memory().percent
    return f"CPU: {cpu}% | RAM: {ram}%"

def list_files(directory_path="."):
    """Liste les fichiers d'un dossier."""
    try:
        if directory_path in ["", ".", "./"]: 
            directory_path = os.getcwd()
        files = os.listdir(directory_path)
        return f"Fichiers dans {directory_path}: " + ", ".join(files)
    except Exception as e:
        return f"Erreur : {str(e)}"

def summarize_directory(directory_path: str = "."):
    """Lit et résume les fichiers d'un dossier (txt, py, pdf, docx)."""
    if directory_path in ["", ".", "./"]:
        directory_path = os.getcwd()

    if not os.path.exists(directory_path):
        return f"Erreur : Le chemin '{directory_path}' n'existe pas."

    combined_text = ""
    files_processed = []

    try:
        for filename in os.listdir(directory_path):
            file_path = os.path.join(directory_path, filename)

            if filename.endswith(('.txt', '.md', '.py')):
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        combined_text += f"\n[Fichier: {filename}]\n{f.read()[:1500]}\n"
                        files_processed.append(filename)
                except: 
                    continue

            elif filename.endswith('.pdf'):
                try:
                    with open(file_path, 'rb') as f:
                        reader = PyPDF2.PdfReader(f)
                        pdf_text = "".join([page.extract_text() or "" for page in reader.pages[:3]])
                        combined_text += f"\n[Fichier PDF: {filename}]\n{pdf_text[:1500]}\n"
                        files_processed.append(filename)
                except: 
                    continue

            elif filename.endswith('.docx'):
                try:
                    doc = Document(file_path)
                    full_text = [para.text for para in doc.paragraphs]
                    docx_content = "\n".join(full_text)[:1500]
                    combined_text += f"\n[Fichier Word: {filename}]\n{docx_content}\n"
                    files_processed.append(filename)
                except: 
                    continue

    except Exception as e:
        return f"Erreur d'accès : {str(e)}"

    if not combined_text:
        return "Aucun document (txt, py, pdf, docx) trouvé."

    prompt = f"Fais un résumé structuré des documents suivants :\n\n{combined_text}"

    try:
        response = ollama.chat(model="llama3.2", messages=[{"role": "user", "content": prompt}])
        return f"Analyse de : {', '.join(files_processed)}\n\n{response['message']['content']}"
    except Exception as e:
        return f"Erreur Ollama : {str(e)}"

def read_specific_file(file_path: str):
    """Lit le contenu d'un fichier spécifique."""
    if not os.path.exists(file_path):
        return f"Erreur : Le fichier '{file_path}' est introuvable."

    try:
        if file_path.endswith(('.txt', '.md', '.py')):
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()[:3000]

        elif file_path.endswith('.pdf'):
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                return "".join([p.extract_text() for p in reader.pages[:5]])[:3000]

        elif file_path.endswith('.docx'):
            doc = Document(file_path)
            return "\n".join([p.text for p in doc.paragraphs])[:3000]

        return "Format de fichier non supporté."
    except Exception as e:
        return f"Erreur lors de la lecture : {str(e)}"

def triage_emails(max_results=5):
    """Trie et affiche les emails importants."""
    return (
        "Liste des emails triés :\n"
        "[Urgent] Facture EDF - À payer avant le 10 février.\n"
        "[Info] Réunion ENSIAS - Présentation IoT lundi prochain."
    )

def manage_tasks(action: str, task_text: str = ""):
    """Actions: 'add' pour sauvegarder, 'list' pour afficher les rappels."""
    filename = "reminders.txt"
    if action == "add":
        with open(filename, "a") as f:
            f.write(f"[{datetime.now().strftime('%Y-%m-%d %H:%M')}] {task_text}\n")
        return f"Rappel sauvegardé : {task_text}"
    elif action == "list":
        if not os.path.exists(filename): 
            return "Aucun rappel trouvé."
        with open(filename, "r") as f:
            return f.read()

# ========================================
# STREAMLIT APP
# ========================================

st.set_page_config(page_title="AI Agent OS Supervisor", page_icon="🤖", layout="wide")

st.title("🤖 Agent TalentAI-Bot")
st.markdown("Interface de supervision et d'analyse de documents (PDF, Word, Texte)")

# ========================================
# SIDEBAR
# ========================================

with st.sidebar:
    # PR #3: File Upload
    st.header("📁 Upload de Documents")
    
    uploaded_files = st.file_uploader(
        "Déposez vos fichiers ici",
        accept_multiple_files=True,
        type=['pdf', 'txt', 'docx', 'doc', 'md', 'py'],
        help="Les fichiers seront analysables immédiatement"
    )
    
    if uploaded_files:
        os.makedirs(TEMP_DIR, exist_ok=True)
        
        saved_files = []
        for uploaded_file in uploaded_files:
            file_path = os.path.join(TEMP_DIR, uploaded_file.name)
            with open(file_path, "wb") as f:
                f.write(uploaded_file.getbuffer())
            saved_files.append(uploaded_file.name)
        
        st.success(f"✅ {len(saved_files)} fichier(s) uploadé(s)")
        with st.expander("Fichiers disponibles"):
            for fname in saved_files:
                st.text(f"📄 {fname}")
    
    st.divider()
    
    # System Stats
    st.header("Système")
    if st.button("Actualiser les stats"):
        stats = get_system_stats()
        st.write(stats)
    
    # PR #2: Conversation Management
    st.divider()
    col1, col2 = st.columns(2)
    with col1:
        msg_count = len([m for m in st.session_state.get('messages', []) if m.get("role") != "system"])
        st.metric("Messages", msg_count)
    with col2:
        if st.button("🗑️ Effacer", help="Réinitialiser la conversation"):
            st.session_state.messages = [st.session_state.messages[0]]
            st.rerun()
    
    st.divider()
    st.info("Agents : Supervisor, Summarizer, Email Triage, Task Manager")

# ========================================
# SESSION STATE INITIALIZATION
# ========================================

if "messages" not in st.session_state:
    st.session_state.messages = [
        {
            "role": "system", 
            "content": (
                "Tu es un assistant IA polyvalent nommé TalentAI-Bot. "
                "Tu peux fonctionner comme un Chatbot ou comme un Agent. "
                f"Les fichiers uploadés se trouvent dans '{TEMP_DIR}/'. "
                "Tu as accès à 'manage_tasks' pour gérer les rappels. "
                "Si l'utilisateur mentionne 'mails', 'emails' ou 'courriels', "
                "tu dois EXCLUSIVEMENT utiliser l'outil 'triage_emails'. "
                "N'utilise jamais 'list_files' ou 'summarize_directory' pour les emails."
            )
        }
    ]

# ========================================
# CHAT DISPLAY
# ========================================

for message in st.session_state.messages:
    if message["role"] != "system":
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

# ========================================
# CHAT INPUT & PROCESSING
# ========================================

if prompt := st.chat_input("Comment puis-je vous aider ?"):
    st.session_state.messages.append({"role": "user", "content": prompt})
    with st.chat_message("user"):
        st.markdown(prompt)

    with st.chat_message("assistant"):
        available_tools = [
            get_system_stats, 
            list_files, 
            summarize_directory, 
            read_specific_file, 
            triage_emails,
            manage_tasks
        ]

        # PR #2: Use pruned messages
        response = ollama.chat(
            model="llama3.2", 
            messages=prune_conversation(st.session_state.messages), 
            tools=available_tools
        )

        if response.get('message', {}).get('tool_calls'):
            st.session_state.messages.append(response['message'])

            # PR #1: Real-time Tool Execution Feedback
            for tool in response['message']['tool_calls']:
                name = tool['function']['name']
                args = tool['function']['arguments']
                
                with st.status(f"🔧 Exécution: {name}...", expanded=True) as status:
                    st.write(f"**Arguments:** `{args}`")
                    
                    if name == "get_system_stats": 
                        result = get_system_stats()
                    elif name == "list_files": 
                        result = list_files(**args)
                    elif name == "summarize_directory": 
                        result = summarize_directory(**args)
                    elif name == "read_specific_file": 
                        result = read_specific_file(**args)
                    elif name == "triage_emails": 
                        result = triage_emails(**args)
                    elif name == "manage_tasks": 
                        result = manage_tasks(**args)
                    else: 
                        result = f"Outil '{name}' inconnu."
                    
                    st.write("**Résultat:**")
                    preview = str(result)[:500] + ("..." if len(str(result)) > 500 else "")
                    st.code(preview, language="text")
                    status.update(label=f"✅ {name} terminé", state="complete")
                
                st.session_state.messages.append({'role': 'tool', 'content': result})

            final_res = ollama.chat(model="llama3.2", messages=prune_conversation(st.session_state.messages))
            full_response = final_res['message']['content']
        else:
            full_response = response['message']['content']

        st.markdown(full_response)
        st.session_state.messages.append({"role": "assistant", "content": full_response})import streamlit as st
import ollama
import os
from datetime import datetime
import psutil
import PyPDF2
from docx import Document

# ========================================
# CONFIGURATION
# ========================================
MAX_MESSAGES = 20
PRESERVE_SYSTEM = True
TEMP_DIR = "temp_uploads"

# ========================================
# UTILITY FUNCTIONS
# ========================================

def prune_conversation(messages, max_messages=MAX_MESSAGES):
    """Keep only recent messages to prevent token overflow"""
    if len(messages) <= max_messages:
        return messages
    
    if PRESERVE_SYSTEM and messages[0]["role"] == "system":
        return [messages[0]] + messages[-(max_messages-1):]
    
    return messages[-max_messages:]

# ========================================
# TOOL FUNCTIONS
# ========================================

def get_system_stats():
    """Returns real-time CPU and RAM usage percentage."""
    cpu = psutil.cpu_percent(interval=1)
    ram = psutil.virtual_memory().percent
    return f"CPU: {cpu}% | RAM: {ram}%"

def list_files(directory_path="."):
    """Liste les fichiers d'un dossier."""
    try:
        if directory_path in ["", ".", "./"]: 
            directory_path = os.getcwd()
        files = os.listdir(directory_path)
        return f"Fichiers dans {directory_path}: " + ", ".join(files)
    except Exception as e:
        return f"Erreur : {str(e)}"

def summarize_directory(directory_path: str = "."):
    """Lit et résume les fichiers d'un dossier (txt, py, pdf, docx)."""
    if directory_path in ["", ".", "./"]:
        directory_path = os.getcwd()

    if not os.path.exists(directory_path):
        return f"Erreur : Le chemin '{directory_path}' n'existe pas."

    combined_text = ""
    files_processed = []

    try:
        for filename in os.listdir(directory_path):
            file_path = os.path.join(directory_path, filename)

            if filename.endswith(('.txt', '.md', '.py')):
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        combined_text += f"\n[Fichier: {filename}]\n{f.read()[:1500]}\n"
                        files_processed.append(filename)
                except: 
                    continue

            elif filename.endswith('.pdf'):
                try:
                    with open(file_path, 'rb') as f:
                        reader = PyPDF2.PdfReader(f)
                        pdf_text = "".join([page.extract_text() or "" for page in reader.pages[:3]])
                        combined_text += f"\n[Fichier PDF: {filename}]\n{pdf_text[:1500]}\n"
                        files_processed.append(filename)
                except: 
                    continue

            elif filename.endswith('.docx'):
                try:
                    doc = Document(file_path)
                    full_text = [para.text for para in doc.paragraphs]
                    docx_content = "\n".join(full_text)[:1500]
                    combined_text += f"\n[Fichier Word: {filename}]\n{docx_content}\n"
                    files_processed.append(filename)
                except: 
                    continue

    except Exception as e:
        return f"Erreur d'accès : {str(e)}"

    if not combined_text:
        return "Aucun document (txt, py, pdf, docx) trouvé."

    prompt = f"Fais un résumé structuré des documents suivants :\n\n{combined_text}"

    try:
        response = ollama.chat(model="llama3.2", messages=[{"role": "user", "content": prompt}])
        return f"Analyse de : {', '.join(files_processed)}\n\n{response['message']['content']}"
    except Exception as e:
        return f"Erreur Ollama : {str(e)}"

def read_specific_file(file_path: str):
    """Lit le contenu d'un fichier spécifique."""
    if not os.path.exists(file_path):
        return f"Erreur : Le fichier '{file_path}' est introuvable."

    try:
        if file_path.endswith(('.txt', '.md', '.py')):
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()[:3000]

        elif file_path.endswith('.pdf'):
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                return "".join([p.extract_text() for p in reader.pages[:5]])[:3000]

        elif file_path.endswith('.docx'):
            doc = Document(file_path)
            return "\n".join([p.text for p in doc.paragraphs])[:3000]

        return "Format de fichier non supporté."
    except Exception as e:
        return f"Erreur lors de la lecture : {str(e)}"

def triage_emails(max_results=5):
    """Trie et affiche les emails importants."""
    return (
        "Liste des emails triés :\n"
        "[Urgent] Facture EDF - À payer avant le 10 février.\n"
        "[Info] Réunion ENSIAS - Présentation IoT lundi prochain."
    )

def manage_tasks(action: str, task_text: str = ""):
    """Actions: 'add' pour sauvegarder, 'list' pour afficher les rappels."""
    filename = "reminders.txt"
    if action == "add":
        with open(filename, "a") as f:
            f.write(f"[{datetime.now().strftime('%Y-%m-%d %H:%M')}] {task_text}\n")
        return f"Rappel sauvegardé : {task_text}"
    elif action == "list":
        if not os.path.exists(filename): 
            return "Aucun rappel trouvé."
        with open(filename, "r") as f:
            return f.read()

# ========================================
# STREAMLIT APP
# ========================================

st.set_page_config(page_title="AI Agent OS Supervisor", page_icon="🤖", layout="wide")

st.title("🤖 Agent TalentAI-Bot")
st.markdown("Interface de supervision et d'analyse de documents (PDF, Word, Texte)")

# ========================================
# SIDEBAR
# ========================================

with st.sidebar:
    # PR #3: File Upload
    st.header("📁 Upload de Documents")
    
    uploaded_files = st.file_uploader(
        "Déposez vos fichiers ici",
        accept_multiple_files=True,
        type=['pdf', 'txt', 'docx', 'doc', 'md', 'py'],
        help="Les fichiers seront analysables immédiatement"
    )
    
    if uploaded_files:
        os.makedirs(TEMP_DIR, exist_ok=True)
        
        saved_files = []
        for uploaded_file in uploaded_files:
            file_path = os.path.join(TEMP_DIR, uploaded_file.name)
            with open(file_path, "wb") as f:
                f.write(uploaded_file.getbuffer())
            saved_files.append(uploaded_file.name)
        
        st.success(f"✅ {len(saved_files)} fichier(s) uploadé(s)")
        with st.expander("Fichiers disponibles"):
            for fname in saved_files:
                st.text(f"📄 {fname}")
    
    st.divider()
    
    # System Stats
    st.header("Système")
    if st.button("Actualiser les stats"):
        stats = get_system_stats()
        st.write(stats)
    
    # PR #2: Conversation Management
    st.divider()
    col1, col2 = st.columns(2)
    with col1:
        msg_count = len([m for m in st.session_state.get('messages', []) if m.get("role") != "system"])
        st.metric("Messages", msg_count)
    with col2:
        if st.button("🗑️ Effacer", help="Réinitialiser la conversation"):
            st.session_state.messages = [st.session_state.messages[0]]
            st.rerun()
    
    st.divider()
    st.info("Agents : Supervisor, Summarizer, Email Triage, Task Manager")

# ========================================
# SESSION STATE INITIALIZATION
# ========================================

if "messages" not in st.session_state:
    st.session_state.messages = [
        {
            "role": "system", 
            "content": (
                "Tu es un assistant IA polyvalent nommé TalentAI-Bot. "
                "Tu peux fonctionner comme un Chatbot ou comme un Agent. "
                f"Les fichiers uploadés se trouvent dans '{TEMP_DIR}/'. "
                "Tu as accès à 'manage_tasks' pour gérer les rappels. "
                "Si l'utilisateur mentionne 'mails', 'emails' ou 'courriels', "
                "tu dois EXCLUSIVEMENT utiliser l'outil 'triage_emails'. "
                "N'utilise jamais 'list_files' ou 'summarize_directory' pour les emails."
            )
        }
    ]

# ========================================
# CHAT DISPLAY
# ========================================

for message in st.session_state.messages:
    if message["role"] != "system":
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

# ========================================
# CHAT INPUT & PROCESSING
# ========================================

if prompt := st.chat_input("Comment puis-je vous aider ?"):
    st.session_state.messages.append({"role": "user", "content": prompt})
    with st.chat_message("user"):
        st.markdown(prompt)

    with st.chat_message("assistant"):
        available_tools = [
            get_system_stats, 
            list_files, 
            summarize_directory, 
            read_specific_file, 
            triage_emails,
            manage_tasks
        ]

        # PR #2: Use pruned messages
        response = ollama.chat(
            model="llama3.2", 
            messages=prune_conversation(st.session_state.messages), 
            tools=available_tools
        )

        if response.get('message', {}).get('tool_calls'):
            st.session_state.messages.append(response['message'])

            # PR #1: Real-time Tool Execution Feedback
            for tool in response['message']['tool_calls']:
                name = tool['function']['name']
                args = tool['function']['arguments']
                
                with st.status(f"🔧 Exécution: {name}...", expanded=True) as status:
                    st.write(f"**Arguments:** `{args}`")
                    
                    if name == "get_system_stats": 
                        result = get_system_stats()
                    elif name == "list_files": 
                        result = list_files(**args)
                    elif name == "summarize_directory": 
                        result = summarize_directory(**args)
                    elif name == "read_specific_file": 
                        result = read_specific_file(**args)
                    elif name == "triage_emails": 
                        result = triage_emails(**args)
                    elif name == "manage_tasks": 
                        result = manage_tasks(**args)
                    else: 
                        result = f"Outil '{name}' inconnu."
                    
                    st.write("**Résultat:**")
                    preview = str(result)[:500] + ("..." if len(str(result)) > 500 else "")
                    st.code(preview, language="text")
                    status.update(label=f"✅ {name} terminé", state="complete")
                
                st.session_state.messages.append({'role': 'tool', 'content': result})

            final_res = ollama.chat(model="llama3.2", messages=prune_conversation(st.session_state.messages))
            full_response = final_res['message']['content']
        else:
            full_response = response['message']['content']

        st.markdown(full_response)
        st.session_state.messages.append({"role": "assistant", "content": full_response})
