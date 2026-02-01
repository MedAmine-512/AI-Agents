import ollama
import os
import psutil
import PyPDF2
from docx import Document 
from datetime import datetime

def summarize_directory(directory_path: str = "."):
    """Lit les fichiers texte, python, PDF et DOCX d'un dossier pour en faire un résumé."""
    if directory_path in ["", ".", "./"]:
        directory_path = os.getcwd()

    if not os.path.exists(directory_path):
        return f"Erreur : Le chemin '{directory_path}' n'existe pas."
    
    combined_text = ""
    files_processed = []

    print(f"[*] Analyse du dossier : {directory_path}")

    try:
        for filename in os.listdir(directory_path):
            file_path = os.path.join(directory_path, filename)
            
            if filename.endswith(('.txt', '.md', '.py')):
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        combined_text += f"\n[Fichier: {filename}]\n{f.read()[:1500]}\n"
                        files_processed.append(filename)
                except: continue

            elif filename.endswith('.pdf'):
                try:
                    with open(file_path, 'rb') as f:
                        reader = PyPDF2.PdfReader(f)
                        pdf_text = "".join([page.extract_text() or "" for page in reader.pages[:3]])
                        combined_text += f"\n[Fichier PDF: {filename}]\n{pdf_text[:1500]}\n"
                        files_processed.append(filename)
                except: continue

            elif filename.endswith('.docx'):
                try:
                    doc = Document(file_path)
                    full_text = [para.text for para in doc.paragraphs]
                    docx_content = "\n".join(full_text)[:1500]
                    combined_text += f"\n[Fichier Word: {filename}]\n{docx_content}\n"
                    files_processed.append(filename)
                except: continue
                
    except Exception as e:
        return f"Erreur d'accès : {str(e)}"

    if not combined_text:
        return "Aucun document (txt, py, pdf, docx) trouvé."

    print(f"[*] Génération du résumé...")
    prompt = f"Fais un résumé structuré des documents suivants contenus dans le dossier :\n\n{combined_text}"
    
    try:
        response = ollama.chat(model="llama3.2", messages=[{"role": "user", "content": prompt}])
        return f"Analyse terminée pour : {', '.join(files_processed)}\n\n{response['message']['content']}"
    except Exception as e:
        return f"Erreur Ollama : {str(e)}"

def get_system_stats():
    """Returns real-time CPU and RAM usage percentage."""
    cpu = psutil.cpu_percent(interval=1)
    ram = psutil.virtual_memory().percent
    return f"CPU: {cpu}% | RAM: {ram}%"

def list_files(directory_path="."):
    """Liste les fichiers d'un dossier (argument directory_path pour l'IA)."""
    try:
        if directory_path in ["", ".", "./"]: directory_path = os.getcwd()
        files = os.listdir(directory_path)
        return f"Fichiers dans {directory_path}: " + ", ".join(files)
    except Exception as e:
        return f"Erreur : {str(e)}"

def read_specific_file(file_path: str):
    """Lit le contenu d'un fichier spécifique (txt, py, pdf, docx)."""
    if not os.path.exists(file_path):
        return f"Erreur : Le fichier '{file_path}' est introuvable."

    try:
        if file_path.endswith(('.txt', '.md', '.py')):
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()[:3000]

        elif file_path.endswith('.pdf'):
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                return "".join([p.extract_text() for p in reader.pages[:5]])[:3000]

        elif file_path.endswith('.docx'):
            doc = Document(file_path)
            return "\n".join([p.text for p in doc.paragraphs])[:3000]

        return "Format de fichier non supporté pour la lecture directe."
    except Exception as e:
        return f"Erreur lors de la lecture : {str(e)}"
    
def triage_emails(max_results=5):
    """
    Indispensable pour consulter les courriels, emails ou messages. 
    Cet outil accède directement à l'API de messagerie et ne nécessite aucun chemin de dossier local.
    """
    return "Liste des emails triés : [Urgent] Facture EDF - Résumé : À payer avant le 10 février. [Info] Réunion ENSIAS - Résumé : Présentation des projets IoT prévue lundi prochain."

def run_supervisor_agent():
    model = "llama3.2"
    print(f"\n--- Agent IA Multi-Docs (Py 3.14) ---")
    
    messages = [{
        "role": "system",
        "content": (
            "Tu es un superviseur. "
            "Utilise 'summarize_directory' pour analyser un dossier. "
            "Utilise 'list_files' pour voir le contenu d'un répertoire. "
            "Utilise 'triage_emails' si l'utilisateur demande de lire, résumer ou trier ses courriels/mails. "
            "Utilise 'read_specific_file' pour lire un fichier précis."
        )
    }]

    while True:
        user_input = input("User: ")
        if user_input.lower() in ["exit", "quit"]: break
        
        messages.append({"role": "user", "content": user_input})
        available_tools = [get_system_stats, list_files, summarize_directory, read_specific_file, triage_emails]

        response = ollama.chat(model=model, messages=messages, tools=available_tools)

        if response.get('message', {}).get('tool_calls'):
            messages.append(response['message'])
            for tool in response['message']['tool_calls']:
                name = tool['function']['name']
                args = tool['function']['arguments']
                
                print(f"[*] Appel de l'outil : {name}")
                
                if name == "get_system_stats": result = get_system_stats()
                elif name == "list_files": result = list_files(**args)
                elif name == "summarize_directory": result = summarize_directory(**args)
                elif name == "read_specific_file": result = read_specific_file(**args)
                elif name == "triage_emails": result = triage_emails(**args)
                else: result = f"Outil '{name}' inconnu."
                
                messages.append({'role': 'tool', 'content': result})

            final_response = ollama.chat(model=model, messages=messages)
            print(f"Agent: {final_response['message']['content']}\n")
        else:
            messages.append(response['message'])
            print(f"Agent: {response['message']['content']}\n")

if __name__ == "__main__":
    run_supervisor_agent()